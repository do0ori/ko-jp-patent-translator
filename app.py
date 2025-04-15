import logging
import tempfile
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from google import genai
from google.genai.errors import ClientError
from PIL import Image
from pydantic import BaseModel

# Gemini API 설정
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])


# 구조화 모델
class TranslationResult(BaseModel):
    translated_text: str


class ImageTranslation(BaseModel):
    original: str
    translated: str


# 로깅 설정
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)


def retry_with_delay(func, *args, max_retries=5, default_delay=10, **kwargs):
    """
    재시도 로직을 처리하는 헬퍼 함수.
    - func: 호출할 함수
    - *args, **kwargs: 함수에 전달할 인자
    - max_retries: 최대 재시도 횟수
    - default_delay: 기본 대기 시간 (초)
    """
    for attempt in range(max_retries):
        try:
            logging.info(
                f"Attempt {attempt + 1}/{max_retries} for function {func.__name__}"
            )
            return func(*args, **kwargs)
        except ClientError as e:
            # RESOURCE_EXHAUSTED 처리
            if e.code == 429 and e.status == "RESOURCE_EXHAUSTED":
                retry_delay = default_delay
                try:
                    # API 응답에서 retryDelay 값 추출
                    retry_info = next(
                        detail
                        for detail in e.details["error"]["details"]
                        if detail["@type"] == "type.googleapis.com/google.rpc.RetryInfo"
                    )
                    retry_delay = int(retry_info["retryDelay"].strip("s"))
                except (StopIteration, KeyError, ValueError) as ex:
                    logging.warning(
                        f"Failed to extract retryDelay from API response. Using default delay. Error: {ex}"
                    )

                logging.warning(
                    f"RESOURCE_EXHAUSTED 발생. {retry_delay}초 후 재시도합니다... (시도 {attempt + 1}/{max_retries})"
                )
                time.sleep(retry_delay)
            else:
                # 다른 ClientError는 즉시 실패
                logging.error(f"Unexpected ClientError in {func.__name__}: {e}")
                raise e
        except Exception as e:
            # 다른 예외는 즉시 실패
            logging.error(f"Unexpected error in {func.__name__}: {e}")
            raise e

    # 최대 재시도 횟수 초과
    logging.error(f"Failed to execute {func.__name__} after {max_retries} retries.")
    raise RuntimeError(
        f"Failed to execute {func.__name__} after {max_retries} retries."
    )


# 텍스트 번역
def translate_text_with_gemini(text: str) -> str:
    def call_gemini_api():
        prompt = (
            "Translate the following Korean patent document text into Japanese. "
            "Translate it naturally, but maintain technical and structural fidelity. "
            "Translate domain-specific technical terms with reference to official or trusted Japanese sources, rather than literal translation."
        )
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=[prompt, text],
            config={
                "response_mime_type": "application/json",
                "response_schema": TranslationResult,
            },
        )
        return response.parsed.translated_text

    return retry_with_delay(call_gemini_api)


# 이미지 내 텍스트 번역
def translate_image_with_gemini(pil_image) -> list[ImageTranslation]:
    def call_gemini_api():
        prompt = (
            "Extract all visible Korean or English text from this image, and translate each into Japanese. "
            "Format the result as a list of objects with 'original' and 'translated' keys."
        )
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=[prompt, pil_image],
            config={
                "response_mime_type": "application/json",
                "response_schema": list[ImageTranslation],
            },
        )
        return response.parsed

    return retry_with_delay(call_gemini_api)


# 초기 상태
if "translated" not in st.session_state:
    st.session_state.translated = False
if "output_path" not in st.session_state:
    st.session_state.output_path = None
if "parsed_elements" not in st.session_state:
    st.session_state.parsed_elements = []
if "chunked_elements" not in st.session_state:
    st.session_state.chunked_elements = []

# 페이지 기본 정보
st.set_page_config(page_title="한일 특허 번역기", layout="centered")
st.title("📄 한일 특허 번역기")
st.markdown(
    "업로드한 특허 문서를 자동 분석하여 Ai 기반으로 한일 번역 문서를 생성합니다."
)

# 파일 업로드
uploaded_file = st.file_uploader("📤 번역할 .docx 파일을 업로드하세요", type=["docx"])

# 진행률 표시 위치 확보
progress_placeholder = st.empty()


# 문서 파싱 함수
def parse_docx_with_images(docx_file):
    doc = Document(docx_file)
    elements = []

    rels = doc.part._rels
    image_map = {}

    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.reltype:
            image_bytes = rel_obj.target_part.blob
            image = Image.open(BytesIO(image_bytes))
            image_map[rel_obj.rId] = image

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            elements.append({"type": "TEXT", "content": text})
        for run in para.runs:
            drawing = run._element.find(
                ".//w:drawing",
                namespaces={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                },
            )
            if drawing is not None:
                blip = drawing.find(
                    ".//a:blip",
                    namespaces={
                        "a": "http://schemas.openxmlformats.org/drawingml/2006/main"
                    },
                )
                if blip is not None:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id in image_map:
                        elements.append(
                            {"type": "FIGURE", "content": image_map[embed_id]}
                        )
    return elements


# 단락들을 묶어서 chunk 생성
def group_paragraphs_to_chunks(elements, max_words=2000):
    chunks, buffer, word_count = [], [], 0
    for elem in elements:
        if elem["type"] == "TEXT":
            words = len(elem["content"].split())
            print(elem["content"].split())
            if word_count + words > max_words and buffer:
                chunks.append({"type": "TEXT", "content": "\n".join(buffer)})
                buffer, word_count = [], 0
            buffer.append(elem["content"])
            word_count += words
        elif elem["type"] == "FIGURE":
            if buffer:
                chunks.append({"type": "TEXT", "content": "\n".join(buffer)})
                buffer, word_count = [], 0
            chunks.append(elem)
    if buffer:
        chunks.append({"type": "TEXT", "content": "\n".join(buffer)})
    return chunks


# 출력용 docx 초기화 함수
def create_japanese_patent_docx():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "MS Mincho"
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    return doc


# 번역 실행
def run_translation():
    doc = create_japanese_patent_docx()
    total = len(st.session_state.chunked_elements)

    for i, chunk in enumerate(st.session_state.chunked_elements):
        if chunk["type"] == "TEXT":
            translated = translate_text_with_gemini(chunk["content"])
            doc.add_paragraph(translated)
            chunk["translated"] = translated
        elif chunk["type"] == "FIGURE":
            translated_pairs = translate_image_with_gemini(chunk["content"])
            formatted = [f"{p.original}: {p.translated}" for p in translated_pairs]
            for line in formatted:
                doc.add_paragraph(line)
            chunk["translated"] = formatted

        progress_placeholder.progress(
            (i + 1) / total, text=f"🔄 번역 중... {i + 1} / {total} 청크 완료"
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc.save(tmp_file.name)
        st.session_state.output_path = tmp_file.name

    st.session_state.translated = True


# 파일 제거 or 변경 시 상태 초기화
if uploaded_file is None:
    st.session_state.translated = False
    st.session_state.output_path = None
    st.session_state.parsed_elements = []
    st.session_state.chunked_elements = []
    st.session_state.base_filename = ""
else:
    new_filename = uploaded_file.name
    if st.session_state.get("last_uploaded_filename") != new_filename:
        st.session_state.translated = False
        st.session_state.output_path = None
        st.session_state.parsed_elements = []
        st.session_state.chunked_elements = []
        st.session_state.last_uploaded_filename = new_filename
        st.session_state.base_filename = Path(new_filename).stem

# 파일 업로드 시 문서 파싱
if uploaded_file and not st.session_state.translated:
    elements = parse_docx_with_images(uploaded_file)
    chunks = group_paragraphs_to_chunks(elements)
    st.session_state.parsed_elements = elements
    st.session_state.chunked_elements = chunks

# 번역 시작 버튼
if uploaded_file and not st.session_state.translated:
    st.button("🚀 번역 시작", on_click=run_translation)

# 번역 완료 후 결과
if st.session_state.translated:
    st.success("✅ 번역이 완료되었습니다!")

    download_filename = f"{st.session_state.base_filename}_translated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    with open(st.session_state.output_path, "rb") as f:
        st.download_button(
            label="📥 번역된 .docx 다운로드",
            data=f,
            file_name=download_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # 번역 결과 표시
    with st.expander("📘 최종 번역 결과 (청크 단위)", expanded=False):
        df = pd.DataFrame(st.session_state.chunked_elements)
        st.dataframe(df, use_container_width=True)
