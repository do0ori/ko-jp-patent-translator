from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image


def parse_docx_with_images(docx_file):
    doc = Document(docx_file)
    elements = []

    rels = doc.part._rels
    image_map = {}

    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.reltype:
            image_bytes = rel_obj.target_part.blob
            image = Image.open(BytesIO(image_bytes))
            image_map[rel_obj.rId] = image

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            elements.append({"type": "TEXT", "content": text})
        for run in para.runs:
            drawing = run._element.find(
                ".//w:drawing",
                namespaces={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                },
            )
            if drawing is not None:
                blip = drawing.find(
                    ".//a:blip",
                    namespaces={
                        "a": "http://schemas.openxmlformats.org/drawingml/2006/main"
                    },
                )
                if blip is not None:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id in image_map:
                        elements.append(
                            {"type": "FIGURE", "content": image_map[embed_id]}
                        )
    return elements


def create_japanese_patent_docx():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "MS Gothic"
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "MS Gothic")
    return doc
